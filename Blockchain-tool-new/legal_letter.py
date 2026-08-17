"""
LEGAL LETTER TEMPLATE GENERATOR

Generates a draft "letter before further action" style document for a
solicitor to send to a KYC exchange, when a trace has found funds
landing at an address attributed to that exchange. This is an OPTIONAL
tool the solicitor chooses to use per finding - nothing in this app
generates or sends this automatically.

IMPORTANT, deliberately: this is a DRAFT for a qualified solicitor to
review, complete, and take responsibility for - not a ready-to-send
legal document, and not legal advice from this tool. It does not
overclaim what a letter alone can compel - a letter cannot force an
exchange to hand over KYC/account-holder data; that typically requires
formal legal process (e.g. a court order). The template asks for
voluntary preservation and cooperation, and flags that further legal
process may follow, rather than assuming compliance.
"""

from datetime import datetime, timezone
from io import BytesIO

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _add_placeholder_paragraph(doc, text, bold=False, size=11, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def generate_legal_letter_docx(exchange_name, wallet_address, chain, amount_label,
                                tx_hash, tx_time_utc, explorer_url, case_reference=None):
    """Returns the raw bytes of a generated .docx letter, pre-filled
    with the specific finding's details. Everything about the sender,
    client, and case beyond the transaction itself is left as a
    bracketed placeholder for the solicitor to complete."""
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Clear draft banner - same principle as the Evidence Pack's own
    # banner: this tool never lets its own output be mistaken for a
    # finished, authoritative document.
    banner = doc.add_paragraph()
    banner_run = banner.add_run(
        "DRAFT — FOR SOLICITOR REVIEW BEFORE SENDING. This letter was auto-drafted from a blockchain "
        "trace finding. It is not legal advice, and a letter alone cannot compel disclosure of "
        "account-holder data - that typically requires formal legal process. Review, complete every "
        "bracketed placeholder, and apply your own professional judgment before this is sent."
    )
    banner_run.bold = True
    banner_run.font.size = Pt(9)
    banner_run.font.color.rgb = RGBColor(0xB0, 0x00, 0x00)
    doc.add_paragraph()

    _add_placeholder_paragraph(doc, "[YOUR FIRM NAME]", bold=True, size=12)
    _add_placeholder_paragraph(doc, "[Firm address]")
    _add_placeholder_paragraph(doc, "[Firm contact details]")
    doc.add_paragraph()

    today = datetime.now(timezone.utc).strftime("%d %B %Y")
    _add_placeholder_paragraph(doc, today)
    doc.add_paragraph()

    _add_placeholder_paragraph(doc, f"{exchange_name} — Compliance / Legal Department", bold=True)
    _add_placeholder_paragraph(doc, "[Exchange address, if known - otherwise send via their published legal/compliance contact channel]")
    doc.add_paragraph()

    _add_placeholder_paragraph(doc, "PRIVATE & CONFIDENTIAL", bold=True)
    _add_placeholder_paragraph(doc, f"Re: Request for Voluntary Assistance — Suspected Proceeds of Fraud{f' ({case_reference})' if case_reference else ''}", bold=True)
    doc.add_paragraph()

    _add_placeholder_paragraph(
        doc,
        "We act on behalf of [Client name] (\"our client\") in connection with the loss of "
        "cryptocurrency assets believed to be the result of [fraud / theft - describe briefly]."
    )
    _add_placeholder_paragraph(
        doc,
        "Blockchain analysis carried out on our client's behalf has identified that funds "
        "traceable to our client were transferred to a wallet address which our investigation "
        "indicates is under your organisation's control. The specific transaction is detailed below."
    )

    # Table of the auto-filled, factual transaction details.
    table = doc.add_table(rows=0, cols=2)
    table.style = "Light Grid Accent 1"
    rows = [
        ("Destination wallet address", wallet_address),
        ("Blockchain", chain or "-"),
        ("Transaction hash", tx_hash or "-"),
        ("Amount", amount_label or "-"),
        ("Date/time (UTC)", tx_time_utc or "-"),
        ("Block explorer reference", explorer_url or "-"),
    ]
    for label, value in rows:
        row_cells = table.add_row().cells
        row_cells[0].width = Inches(2.1)
        row_cells[1].width = Inches(4.2)
        row_cells[0].paragraphs[0].add_run(label).bold = True
        row_cells[1].paragraphs[0].add_run(str(value))
    doc.add_paragraph()

    _add_placeholder_paragraph(
        doc,
        "We would be grateful if you would, as a matter of urgency and pending any formal legal "
        "process that may follow:"
    )
    for bullet in [
        "Preserve all records relating to the account(s) associated with the above address, "
        "including KYC/onboarding records, transaction history, and any related account activity;",
        "Consider a voluntary freeze on the relevant account pending further investigation, to the "
        "extent your policies and applicable law permit;",
        "Confirm receipt of this letter and provide a reference or contact point for further "
        "correspondence.",
    ]:
        p = doc.add_paragraph(bullet, style="List Bullet")
        p.paragraph_format.space_after = Pt(4)

    _add_placeholder_paragraph(
        doc,
        "We note that formal disclosure of account-holder information typically requires an "
        "appropriate legal basis (for example, a court order). We may in due course make a formal "
        "application for such an order if this matter is not resolved through voluntary cooperation, "
        "and would ask that the records above be preserved in the meantime."
    )
    _add_placeholder_paragraph(
        doc,
        "Given the ongoing nature of this matter, we would be grateful for your urgent attention. "
        "Please do not hesitate to contact us using the details above should you require any further "
        "information."
    )
    doc.add_paragraph()
    _add_placeholder_paragraph(doc, "Yours faithfully,")
    doc.add_paragraph()
    doc.add_paragraph()
    _add_placeholder_paragraph(doc, "[Solicitor name]", bold=True)
    _add_placeholder_paragraph(doc, "[Position]")
    _add_placeholder_paragraph(doc, "[Firm name]")

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
